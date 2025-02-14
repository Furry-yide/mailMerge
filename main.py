
# 开发人员:一德 
# 联系方式:1028647581@qq.com
# 代码版本:v1.0.0
# 本代码仅供学习参考，不得用于商业用途!
# 遵循CC BY-NC 开源协议
# 创建时间:2025年02月14日


import os
from plistlib import InvalidFileException
from docx import Document
from openpyxl import load_workbook

class FileReadError(Exception):
    """自定义异常类,用于读取Excel文件时抛出异常"""
    pass

def readExcel(dataBasePath="data/dataBase.xlsx"):
    """
    读取Excel文件\n
    并读取里面全部的字段\n
    `filePath` => 数据文件路径\n
    `return` => 返回一个元组，第一个元素是表头，第二个元素是数据\n
    未定义 `filePath` 时默认读取 `data/database.xlsx` 文件
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(dataBasePath):
            raise FileNotFoundError(f"文件路径 {dataBasePath} 不存在。")

        wb = load_workbook(dataBasePath)
        ws = wb.active
        headers = [cell.value for cell in ws[1]]
        data = [row for row in ws.iter_rows(min_row=2, values_only=True)]  # 假设第一行是标题

        return headers, data
    
    except FileNotFoundError as e:
        raise(FileReadError(f"文件未找到: {e}"))
    except InvalidFileException as e:
        raise FileReadError(f"无效的Excel文件: {e}")
    except Exception as e:
        raise FileReadError(f"读取Excel文件时发生错误: {e}")



def buildWord(buildPath="build/",buildName= 'file', templatePath="data/template.docx", excelHeader=[], data=[], type=False):
    """
    替换模版中的关键字\n
    并生成新的Word文件\n

    `templatePath` => 模版文件路径\n
    `excelHeader` => Excel表头\n
    `data` => Excel数据\n
    `buildPath` => 生成文件路径\n
    `buildName` => 生成文件名\n
    `type` => 是否生成多个文件,当 False 将生成一个文件,当 True 为生成多个文件\n
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(templatePath):
            raise FileNotFoundError(f"文件路径 {templatePath} 不存在。")

        # 创建或确保 buildPath 文件夹存在
        os.makedirs(buildPath, exist_ok=True)

        # 生成单个文件
        if type != True:
            # 创建一个新的文档用于合并所有内容
            merged_doc = Document()

            for idx, row in enumerate(data):
                # 创建或复制模板
                doc = Document(templatePath)
                for i in range(len(excelHeader)):
                    student = row[i] 
                    # 替换关键字
                    for paragraph in doc.paragraphs:
                        if f"{{{excelHeader[i]}}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace(f"{{{excelHeader[i]}}}", f"{student}")

                # 将当前文档的内容添加到合并文档中
                for element in doc.element.body:
                    merged_doc.element.body.append(element)

                # # 如果不是最后一个学生，添加分页符
                # if idx < len(data) - 1:
                #     merged_doc.add_page_break()

            # 保存合并后的文档
            merged_path = f"{buildPath}{buildName}.docx"
            merged_doc.save(merged_path)
        

        # 生成多个文件
        if type != False:
            for idx, row in enumerate(data):
                # 创建或复制模板
                doc = Document(templatePath)
                for i in range(len(excelHeader)):
                    student = row[i] 
                    # 替换关键字
                    for paragraph in doc.paragraphs:
                        if f"{{{excelHeader[i]}}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace(f"{{{excelHeader[i]}}}", f"{student}")

                        # 保存生成的文档
                        doc.save(rf"{buildPath}{buildName}{idx + 1}.docx")
    
    except FileNotFoundError as e:
        raise(FileReadError(f"文件未找到: {e}"))
    except InvalidFileException as e:
        raise FileReadError(f"无效的Word文件: {e}")
    except Exception as e:
        raise FileReadError(f"读取Word文件时发生错误: {e}")


# readExcel可以传参 dataBasePath="data/dataBase.xlsx" 数据源路径
headerData , studentData = readExcel() # 读取Excel文件

# buildWord可以传参 templatePath="data/template.docx" 模版路径
buildWord(excelHeader=headerData, data=studentData) # 生成Word文件